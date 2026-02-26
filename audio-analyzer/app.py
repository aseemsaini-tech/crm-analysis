"""
Audio Transcription & Analysis App
- Transcribes audio using AssemblyAI
- Analyzes transcript using Claude (Anthropic API)
- Exports transcript as DOCX
- Exports analyzed attributes as CSV
"""

import os
import json
import csv
import time
import tempfile
import secrets
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, jsonify, send_file
import assemblyai as aai
import anthropic
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 500 * 1024 * 1024  # 500MB max upload
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", secrets.token_hex(32))

# Use /tmp on Render (ephemeral filesystem)
BASE_TMP = Path(tempfile.gettempdir()) / "audio-analyzer"
UPLOAD_DIR = BASE_TMP / "uploads"
OUTPUT_DIR = BASE_TMP / "outputs"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ─── Store session data in-memory ───
session_store = {}


def get_clients():
    """Initialize API clients from environment variables."""
    assemblyai_key = os.environ.get("ASSEMBLYAI_API_KEY", "")
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY", "")

    if not assemblyai_key:
        raise ValueError("ASSEMBLYAI_API_KEY environment variable is not set")
    if not anthropic_key:
        raise ValueError("ANTHROPIC_API_KEY environment variable is not set")

    aai.settings.api_key = assemblyai_key
    claude_client = anthropic.Anthropic(api_key=anthropic_key)
    return claude_client


# ─── Routes ──────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/transcribe", methods=["POST"])
def transcribe():
    """Upload audio and transcribe with AssemblyAI."""
    try:
        get_clients()
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    if "audio" not in request.files:
        return jsonify({"error": "No audio file provided"}), 400

    audio_file = request.files["audio"]
    if audio_file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    # Save uploaded file
    filename = f"{int(time.time())}_{audio_file.filename}"
    filepath = UPLOAD_DIR / filename
    audio_file.save(filepath)

    try:
        # Configure transcription
        config = aai.TranscriptionConfig(
            speaker_labels=True,
            auto_highlights=True,
        )
        transcriber = aai.Transcriber()
        transcript = transcriber.transcribe(str(filepath), config=config)

        if transcript.status == aai.TranscriptStatus.error:
            return jsonify({"error": f"Transcription failed: {transcript.error}"}), 500

        # Build structured transcript
        utterances = []
        if transcript.utterances:
            for u in transcript.utterances:
                utterances.append({
                    "speaker": u.speaker,
                    "text": u.text,
                    "start": u.start,
                    "end": u.end,
                })

        full_text = transcript.text or ""

        # Store in session
        session_id = filename.split(".")[0]
        session_store[session_id] = {
            "filename": audio_file.filename,
            "full_text": full_text,
            "utterances": utterances,
            "duration_seconds": (transcript.audio_duration or 0),
            "timestamp": datetime.now().isoformat(),
        }

        return jsonify({
            "session_id": session_id,
            "text": full_text,
            "utterances": utterances,
            "duration_seconds": transcript.audio_duration,
            "word_count": len(full_text.split()),
        })

    except Exception as e:
        return jsonify({"error": f"Transcription error: {str(e)}"}), 500
    finally:
        # Cleanup uploaded file
        if filepath.exists():
            filepath.unlink()


@app.route("/api/analyze", methods=["POST"])
def analyze():
    """Analyze transcript with Claude using a custom prompt."""
    try:
        claude_client = get_clients()
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    data = request.get_json()
    session_id = data.get("session_id")
    analysis_prompt = data.get("prompt", "")
    model = data.get("model", "claude-sonnet-4-5-20250929")

    if not session_id or session_id not in session_store:
        return jsonify({"error": "Invalid session. Please transcribe audio first."}), 400
    if not analysis_prompt.strip():
        return jsonify({"error": "Analysis prompt cannot be empty."}), 400

    session = session_store[session_id]
    transcript_text = session["full_text"]

    # Build the message for Claude
    system_prompt = """You are an expert audio transcript analyst.
You will be given a transcript and an analysis prompt.
Respond with a JSON object containing two keys:
1. "summary": A free-text analysis based on the prompt.
2. "attributes": An object with key-value pairs of extracted attributes.
   Keys should be short column-friendly names (snake_case).
   Values should be strings or numbers.

Respond ONLY with valid JSON, no markdown fences."""

    user_message = f"""## Transcript
{transcript_text}

## Analysis Instructions
{analysis_prompt}

Respond with JSON containing "summary" and "attributes" keys."""

    try:
        message = claude_client.messages.create(
            model=model,
            max_tokens=4096,
            system=system_prompt,
            messages=[{"role": "user", "content": user_message}],
        )

        response_text = message.content[0].text.strip()

        # Parse JSON response (handle markdown fences if present)
        if response_text.startswith("```"):
            response_text = response_text.split("\n", 1)[1]
            response_text = response_text.rsplit("```", 1)[0]

        result = json.loads(response_text)
        summary = result.get("summary", "")
        attributes = result.get("attributes", {})

        # Store analysis in session
        session_store[session_id]["analysis"] = {
            "summary": summary,
            "attributes": attributes,
            "prompt_used": analysis_prompt,
            "model": model,
        }

        return jsonify({
            "summary": summary,
            "attributes": attributes,
        })

    except json.JSONDecodeError:
        # If Claude didn't return valid JSON, return the raw text
        session_store[session_id]["analysis"] = {
            "summary": response_text,
            "attributes": {},
            "prompt_used": analysis_prompt,
            "model": model,
        }
        return jsonify({
            "summary": response_text,
            "attributes": {},
            "note": "Claude did not return structured JSON. Try refining your prompt."
        })
    except Exception as e:
        return jsonify({"error": f"Analysis error: {str(e)}"}), 500


@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    """Export transcript as a formatted DOCX file."""
    data = request.get_json()
    session_id = data.get("session_id")

    if not session_id or session_id not in session_store:
        return jsonify({"error": "Invalid session."}), 400

    session = session_store[session_id]

    doc = Document()

    # --- Title ---
    title = doc.add_heading("Audio Transcript", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Metadata ---
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"File: {session['filename']}  |  Date: {session['timestamp'][:10]}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    duration = session.get("duration_seconds", 0)
    if duration:
        minutes = int(duration // 60)
        seconds = int(duration % 60)
        dur_run = meta.add_run(f"  |  Duration: {minutes}m {seconds}s")
        dur_run.font.size = Pt(10)
        dur_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("")  # spacer

    # --- Transcript body ---
    doc.add_heading("Transcript", level=1)

    utterances = session.get("utterances", [])
    if utterances:
        for u in utterances:
            p = doc.add_paragraph()
            speaker_run = p.add_run(f"Speaker {u['speaker']}: ")
            speaker_run.bold = True
            speaker_run.font.size = Pt(11)
            speaker_run.font.color.rgb = RGBColor(44, 62, 80)
            text_run = p.add_run(u["text"])
            text_run.font.size = Pt(11)
    else:
        p = doc.add_paragraph(session["full_text"])
        p.style.font.size = Pt(11)

    # --- Analysis section (if available) ---
    analysis = session.get("analysis")
    if analysis:
        doc.add_page_break()
        doc.add_heading("Analysis", level=1)

        prompt_p = doc.add_paragraph()
        prompt_label = prompt_p.add_run("Prompt used: ")
        prompt_label.bold = True
        prompt_label.font.size = Pt(10)
        prompt_text = prompt_p.add_run(analysis["prompt_used"])
        prompt_text.font.size = Pt(10)
        prompt_text.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph("")
        doc.add_heading("Summary", level=2)
        summary_p = doc.add_paragraph(analysis["summary"])
        summary_p.style.font.size = Pt(11)

        if analysis.get("attributes"):
            doc.add_heading("Extracted Attributes", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Attribute"
            hdr[1].text = "Value"
            for key, val in analysis["attributes"].items():
                row = table.add_row().cells
                row[0].text = str(key)
                row[1].text = str(val)

    # Save
    output_path = OUTPUT_DIR / f"transcript_{session_id}.docx"
    doc.save(str(output_path))

    return send_file(
        str(output_path),
        as_attachment=True,
        download_name=f"transcript_{session['filename'].rsplit('.', 1)[0]}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/api/export/csv", methods=["POST"])
def export_csv():
    """Export analyzed attributes as a CSV row."""
    data = request.get_json()
    session_id = data.get("session_id")

    if not session_id or session_id not in session_store:
        return jsonify({"error": "Invalid session."}), 400

    session = session_store[session_id]
    analysis = session.get("analysis")

    if not analysis or not analysis.get("attributes"):
        return jsonify({"error": "No analysis attributes found. Run analysis first."}), 400

    attributes = analysis["attributes"]

    # Add metadata columns
    row = {
        "filename": session["filename"],
        "date": session["timestamp"][:10],
        "duration_seconds": session.get("duration_seconds", ""),
        "word_count": len(session["full_text"].split()),
        **attributes,
    }

    output_path = OUTPUT_DIR / f"analysis_{session_id}.csv"
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=row.keys())
        writer.writeheader()
        writer.writerow(row)

    return send_file(
        str(output_path),
        as_attachment=True,
        download_name=f"analysis_{session['filename'].rsplit('.', 1)[0]}.csv",
        mimetype="text/csv",
    )


@app.route("/api/sessions", methods=["GET"])
def list_sessions():
    """List active sessions."""
    sessions = []
    for sid, data in session_store.items():
        sessions.append({
            "session_id": sid,
            "filename": data["filename"],
            "timestamp": data["timestamp"],
            "has_analysis": "analysis" in data,
        })
    return jsonify(sessions)


@app.route("/api/health", methods=["GET"])
def health():
    """Health check endpoint for Render."""
    has_aai = bool(os.environ.get("ASSEMBLYAI_API_KEY"))
    has_ant = bool(os.environ.get("ANTHROPIC_API_KEY"))
    return jsonify({
        "status": "ok",
        "assemblyai_configured": has_aai,
        "anthropic_configured": has_ant,
    })


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_ENV") != "production"
    print(f"\n  Audio Transcription & Analysis App")
    print(f"  Running on http://0.0.0.0:{port}")
    print(f"  Debug: {debug}\n")
    app.run(debug=debug, host="0.0.0.0", port=port)
